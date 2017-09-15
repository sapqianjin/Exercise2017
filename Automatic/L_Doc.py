#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Dorian Wang 2017.09.13
# lesson: word file processing

import docx


# TODO: create batch change update date and version for blueprint

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return ('\n'.join(fullText))


# nameString = './Automatic/demo.docx'
nameString = 'demo.docx'
print(getText(nameString))

newDoc = docx.Document()
newDoc.add_paragraph('This is the first page.', 'Title')
newDoc.add_paragraph('Second line.')
newDoc.paragraphs[0].add_run('First line.')
newDoc.paragraphs[0].runs[1].add_break()
newDoc.add_page_break()
newDoc.add_paragraph('This is the second page.')

newDoc.add_heading('Header 0', 0)
newDoc.add_heading('Header 1', 1)
newDoc.add_heading('Header 2', 2)
newDoc.add_heading('Header 3', 3)
newDoc.add_heading('Header 4', 4)

newDoc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))

newDoc.save('DorianDocTest.docx')


dt1=docx.Document('DT_F1.docx')
header = section.header
header.text = 'foobar'

print('test1')
print('test2')